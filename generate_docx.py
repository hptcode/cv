"""Generate Patrick Ho's resume as a two-page Word document (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x29, 0x37)
ACCENT = RGBColor(0x0E, 0x4F, 0x8A)  # deep blue accent for headings
GREY = RGBColor(0x55, 0x5F, 0x6E)


def set_cell_margins(table, top=0, start=0, bottom=0, end=0):
    pass  # not used


def style_paragraph(p, space_before=0, space_after=4, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)


def add_heading(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=8, space_after=3, line=20)
    run = p.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '0E4F8A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_job(doc, title, company, years, points):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=5, space_after=1, line=19)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Calibri'
    r2 = p.add_run('  —  ' + company)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = GREY
    r3 = p.add_run('   (' + years + ')')
    r3.font.size = Pt(10)
    r3.font.name = 'Calibri'
    r3.font.color.rgb = GREY

    for pt in points:
        bp = doc.add_paragraph(style='List Bullet')
        style_paragraph(bp, space_after=1, line=18)
        run = bp.add_run(pt)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # ---- Header ----
    p = doc.add_paragraph()
    style_paragraph(p, space_after=0, line=30)
    r = p.add_run('Patrick Ho')
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    style_paragraph(p, space_after=6)
    r = p.add_run('Systems Support Specialist')
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT
    r = p.add_run('    ·    ho_p@yahoo.com    ·    Vancouver, BC    ·    github.com/hptcode')
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    # ---- Summary ----
    add_heading(doc, 'Career Summary')
    for para in [
        'A systems specialist with over 20 years keeping enterprise IT running—UNIX, Linux, and Windows environments, networking and storage, and Oracle engineered systems. I have worked across systems integration, infrastructure projects, and mission-critical environments, and I enjoy solving hard problems, writing clear documentation, and helping others get up to speed.',
        'More recently I have focused on modern DevOps and self-hosting, deploying my own containerized apps with Docker and Coolify, experimenting with AI developer tooling like Claude Code, and maintaining open-source projects on GitHub.',
    ]:
        sp = doc.add_paragraph()
        style_paragraph(sp, space_after=4, line=17)
        run = sp.add_run(para)
        run.font.size = Pt(10)

    # ---- Core Skills (compact 2-column table) ----
    add_heading(doc, 'Core Skills')
    groups = [
        ('Systems & Platforms', ['UNIX/Linux Administration', 'Windows Server', 'Shell Scripting']),
        ('Network & Infrastructure', ['Cisco Multi-site Networks', 'SAN/NAS Storage', 'Telecom CDR Processing']),
        ('Virtualization & Containers', ['Docker Containerization', 'Virtualization (VMware, KVM, Oracle VM)', 'Hypervisor Management', 'CI/CD Pipelines']),
        ('Cloud & Databases', ['Oracle Cloud (OCI)', 'Oracle RAC Clusters', 'Oracle Engineered Systems', 'Database Support (Postgres, Oracle)']),
        ('Self-Hosted & Dev Tools', ['Self-Hosted Deployment & Coolify (PaaS)', 'AI-Powered Dev Tools & Agents (Claude Code, Hermes)', 'Open-Source Development & GitHub']),
        ('Professional Skills', ['Customer-Facing Support', 'Problem Diagnosis & Resolution', 'Change Management']),
    ]
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    col_widths = [Inches(3.55), Inches(3.55)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
    # Fill cells: groups laid top-to-bottom then left-to-right isn't ideal;
    # do a natural reading order: left column groups 0,2,4 ; right column 1,3,5
    layout = [(0, 0, 0), (1, 1, 0), (2, 0, 1), (3, 1, 1), (4, 0, 2), (5, 1, 2)]
    for gi, col, row in layout:
        label, items = groups[gi]
        cell = table.cell(row, col)
        cell.paragraphs[0].text = ''
        lp = cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = DARK
        lp.paragraph_format.space_after = Pt(1)
        for item in items:
            ip = cell.add_paragraph()
            ip.paragraph_format.space_after = Pt(0)
            ip.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            ip.paragraph_format.line_spacing = Pt(14)
            ir = ip.add_run('•  ' + item)
            ir.font.size = Pt(9)
            ir.font.color.rgb = GREY
    # spacing paragraph after table
    sp = doc.add_paragraph()
    style_paragraph(sp, space_after=2)

    # ---- Key Achievements ----
    add_heading(doc, 'Key Achievements')
    for item in [
        'Deployed Oracle RAC clusters for the 2010 Vancouver Winter Olympics, delivering database infrastructure that stayed reliable under real-time, high-load conditions during an international event.',
        'Designed a national inventory monitoring system at Sun Microsystems—a real-time hardware tracker used by every field engineer across Canada, cutting response times and improving visibility.',
        'Supported mission-critical SAP infrastructure for ICBC (2003–2008), one of BC’s largest enterprises, on Sun environments.',
        'Automated telecom billing at Navigator-Westel, processing thousands of daily CDR records into structured databases for enterprise billing and data warehousing.',
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        style_paragraph(bp, space_after=3, line=17)
        run = bp.add_run(item)
        run.font.size = Pt(10)

    # ---- Experience ----
    add_heading(doc, 'Professional Experience')
    jobs = [
        ('Systems Support Specialist', 'Oracle Canada', '2010 – 2026', [
            'Integrated Oracle engineered systems with Oracle Cloud Infrastructure (OCI), standing up seamless hybrid cloud deployments for customers.',
            'Led OS, SAN, NAS, and backup install and migration projects for enterprise customers, keeping downtime minimal and data intact.',
            'Triaged and resolved issues across Oracle systems, sustaining 99.9%+ uptime for mission-critical environments.',
            'Deployed Oracle RAC cluster infrastructure for the 2010 Vancouver Winter Olympics, keeping databases available in real time under heavy load.',
            'Wrote technical documentation and ran training sessions so customers and partners resolved more issues on their own.',
        ]),
        ('Systems Support Engineer — Global Services', 'Sun Microsystems', '1999 – 2010', [
            'Resolved complex hardware, software, and infrastructure issues for enterprise customers across Canada as a second level escalation resource.',
            'Delivered server and storage infrastructure projects for crown corporations, telecommunication companies, retail chains like Best Buy, and many small and mid-sized businesses.',
            'Provided onsite enterprise UNIX support for ICBC (2003–2008), keeping mission-critical SAP applications on Solaris running smoothly.',
            'Designed and maintained a national inventory monitoring system giving every Canadian field engineer real-time hardware tracking and reporting.',
            'Trained and mentored support teams and wrote documentation that raised the team’s overall capability and knowledge.',
        ]),
        ('Senior Systems & Network Administrator', 'Navigator-Westel Communications', '1997 – 1999', [
            'Managed Cisco multi-site infrastructure for 100+ users across three remote locations, keeping connectivity and services reliable.',
            'Administered 20+ enterprise systems—DNS, mail, Sybase databases, Remedy ITSM, and authentication—across Solaris, Linux, Novell, and Windows.',
            'Automated the processing of thousands of daily telecom CDR records into structured databases for enterprise billing and data warehousing.',
            'Ran helpdesk operations and coordinated infrastructure relocations, keeping services continuous through every transition.',
        ]),
    ]
    for title, company, years, points in jobs:
        add_job(doc, title, company, years, points)

    # ---- Open Source ----
    add_heading(doc, 'Open Source & Projects')
    p = doc.add_paragraph()
    style_paragraph(p, space_after=1)
    r = p.add_run('homeXpensify')
    r.bold = True
    r.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    style_paragraph(p2, space_after=1, line=17)
    r = p2.add_run(
        'A self-hosted, open-source home expense tracker built for multiple users. '
        'A hands-on way I keep current with web development, multi-tenant data modeling, and Docker deployment.'
    )
    r.font.size = Pt(10)
    p3 = doc.add_paragraph()
    style_paragraph(p3, space_after=4)
    r = p3.add_run('Web Development  ·  Multitenant  ·  Docker  ·  Self-Hosted')
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY

    # ---- Education ----
    add_heading(doc, 'Education')
    p = doc.add_paragraph()
    style_paragraph(p, space_after=0)
    r = p.add_run('Bachelor of Science in Electrical Engineering')
    r.bold = True
    r.font.size = Pt(10.5)
    p = doc.add_paragraph()
    style_paragraph(p, space_after=4)
    r = p.add_run('University of Alberta')
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    # ---- Additional ----
    add_heading(doc, 'Additional Information')
    for item in [
        'Languages: English, Cantonese, Mandarin',
        'Location: Vancouver, British Columbia (open to remote/hybrid opportunities)',
        'References available upon request',
    ]:
        bp = doc.add_paragraph(style='List Bullet')
        style_paragraph(bp, space_after=2, line=17)
        run = bp.add_run(item)
        run.font.size = Pt(10)

    out = r'C:\Users\whist\ai-projects\resume\cv-site\Patrick_Ho_Resume_v3.docx'
    doc.save(out)
    print('Saved:', out)


if __name__ == '__main__':
    main()
